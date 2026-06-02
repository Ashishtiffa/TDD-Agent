import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

_ASSETS_DIR = Path(__file__).resolve().parent / 'assets'
HITACHI_LOGO_PATH = _ASSETS_DIR / 'hitachi_solutions_logo.png'
HEADER_BAND_PATH = _ASSETS_DIR / 'header_band.png'


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _style_run(run, size=10, bold=False, color_rgb=None, font_name='Hitachi Sans'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    # Specifically for fonts with spaces or custom fonts in python-docx
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_header_row(table, headers, bg_color="BDD7EE"):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        set_cell_bg(cell, bg_color)
        set_cell_borders(cell)
        for para in cell.paragraphs:
            for run in para.runs:
                _style_run(run, bold=True)


def add_data_row(table, values, border_color="CCCCCC"):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val) if val else ''
        set_cell_borders(cell, border_color)
        for para in cell.paragraphs:
            for run in para.runs:
                _style_run(run)


def add_ax_subsection_table(doc, heading, headers, rows):
    """Simple black-bordered table for AX table TDD subsections."""
    if not rows:
        return
    p = doc.add_paragraph()
    run = p.add_run(heading)
    _style_run(run, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        set_cell_bg(cell, 'BDD7EE')
        set_cell_borders(cell, '000000')
        for para in cell.paragraphs:
            for r in para.runs:
                _style_run(r, bold=True)

    for row_vals in rows:
        add_data_row(table, row_vals, border_color='000000')
    doc.add_paragraph()


def add_ax_table_section(doc, tbl):
    """One AX table block in Data Dictionary — only subsections with data."""
    add_hyperlink_heading(doc, tbl.get('name', 'Table'))

    description = tbl.get('description', '')
    if description:
        add_bold_label(doc, 'Description:', description)

    fields = tbl.get('fields') or []
    if fields:
        add_ax_subsection_table(
            doc,
            'Fields:',
            ['Field Name', 'Type', 'EDT'],
            [[f.get('name', ''), f.get('type', ''), f.get('edt', '')] for f in fields],
        )

    field_groups = tbl.get('field_groups') or []
    if field_groups:
        add_ax_subsection_table(
            doc,
            'Fields Group:',
            ['Fields Group Name', 'Fields'],
            [[g.get('name', ''), g.get('fields', '')] for g in field_groups],
        )

    indexes = tbl.get('indexes') or []
    if indexes:
        add_ax_subsection_table(
            doc,
            'Indexes:',
            ['Index Name', 'Fields'],
            [[idx.get('name', ''), idx.get('fields', '')] for idx in indexes],
        )

    relations = tbl.get('relations') or []
    if relations:
        add_ax_subsection_table(
            doc,
            'Relations:',
            ['Relation name', 'Field', 'Related Table', 'Related Table Field'],
            [[
                r.get('name', ''),
                r.get('field', ''),
                r.get('related_table', ''),
                r.get('related_table_field', ''),
            ] for r in relations],
        )

    methods = tbl.get('methods') or []
    if methods:
        add_ax_subsection_table(
            doc,
            'Methods:',
            ['Method Name', 'Description'],
            [[m.get('name', ''), m.get('description', '')] for m in methods],
        )


def add_section_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_heading(text, level=1)
        _style_run(p.runs[0], size=14, color_rgb=RGBColor(0x1F, 0x49, 0x7D))
    elif level == 2:
        p = doc.add_heading(text, level=2)
        _style_run(p.runs[0], size=12, color_rgb=RGBColor(0x1F, 0x49, 0x7D))
    else:
        p = doc.add_heading(text, level=3)
        _style_run(p.runs[0], size=11, color_rgb=RGBColor(0x1F, 0x49, 0x7D))


def add_hyperlink_heading(doc, text, level=3):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _style_run(run, size=11, bold=True, color_rgb=RGBColor(0x1F, 0x49, 0x7D))
    run.font.underline = True
    return p


def add_bold_label(doc, label, value=''):
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    _style_run(run_label, bold=True)
    if value:
        run_val = p.add_run(' ' + value)
        _style_run(run_val)
    return p


def format_version(version):
    if not version:
        return ''
    v = str(version).strip()
    if v.lower().startswith('dynamics 365'):
        return v
    return f'Dynamics 365 - {v}'


def _style_cell_text(cell, size=11, bold=False, color_rgb=None):
    for para in cell.paragraphs:
        for run in para.runs:
            _style_run(run, size=size, bold=bold, color_rgb=color_rgb)


def add_cover_title_bar(doc):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_bg(cell, 'D9D9D9')
    set_cell_borders(cell, 'D9D9D9')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Technical Design Document')
    _style_run(run, size=16, bold=True, color_rgb=RGBColor(0, 0, 0))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_cover_enhancement_line(doc, header):
    wi = header.get('work_item', '')
    title = header.get('enhancement_title', '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_link = p.add_run(f'Enhancement {wi}' if wi else 'Enhancement')
    _style_run(run_link, size=11, color_rgb=RGBColor(0x05, 0x63, 0xC1))
    run_link.underline = True
    if title:
        run_title = p.add_run(f': {title}')
        _style_run(run_title, size=11, color_rgb=RGBColor(0, 0, 0))
    p.paragraph_format.space_after = Pt(14)


def add_cover_meta_field(doc, label, value, space_after=10):
    p_label = doc.add_paragraph()
    run_label = p_label.add_run(label)
    _style_run(run_label, size=11, color_rgb=RGBColor(0, 0, 0))
    p_label.paragraph_format.space_after = Pt(2)

    p_value = doc.add_paragraph()
    run_value = p_value.add_run(value or '')
    _style_run(run_value, size=11, bold=True, color_rgb=RGBColor(0, 0, 0))
    p_value.paragraph_format.space_after = Pt(space_after)


def add_cover_table_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _style_run(run, size=11, bold=True, color_rgb=RGBColor(0, 0, 0))
    p.paragraph_format.space_after = Pt(4)


def add_cover_black_table(doc, headers, data_rows, revision_data_row=False):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header_text
        set_cell_bg(cell, '000000')
        set_cell_borders(cell, '000000')
        _style_cell_text(cell, size=11, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))

    for row_idx, row_values in enumerate(data_rows):
        row = table.rows[row_idx + 1]
        fill = 'DDEBF7' if revision_data_row else 'FFFFFF'
        for col_idx, value in enumerate(row_values):
            cell = row.cells[col_idx]
            cell.text = str(value) if value else ''
            set_cell_bg(cell, fill)
            set_cell_borders(cell, '000000')
            _style_cell_text(cell, size=11, bold=False, color_rgb=RGBColor(0, 0, 0))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(14)


def _add_paragraph_bottom_border(paragraph, color='808080'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_header(doc):
    """Hitachi Solutions logo + grey line on every page header."""
    logo_path = HITACHI_LOGO_PATH if HITACHI_LOGO_PATH.exists() else HEADER_BAND_PATH
    if not logo_path.exists():
        return

    for section in doc.sections:
        # Distance from top of page to header (larger = logo sits lower, avoids top clip)
        section.header_distance = Inches(0.55)
        header = section.header
        header.is_linked_to_previous = False

        logo_para = header.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_para.paragraph_format.space_before = Pt(0)
        logo_para.paragraph_format.space_after = Pt(4)
        run = logo_para.add_run()
        # Width keeps aspect ratio; slightly smaller than full band to fit header
        run.add_picture(str(logo_path), width=Inches(2.35))

        line_para = header.add_paragraph()
        line_para.paragraph_format.space_before = Pt(2)
        line_para.paragraph_format.space_after = Pt(8)
        _add_paragraph_bottom_border(line_para)


def add_cover_page(doc, header):
    client = header.get('client_name') or header.get('authored_org') or 'Hitachi'
    authored_org = header.get('authored_org') or client
    version = format_version(header.get('version', ''))

    add_cover_title_bar(doc)
    add_cover_enhancement_line(doc, header)

    p_client = doc.add_paragraph()
    run_client = p_client.add_run(client)
    _style_run(run_client, size=11, bold=True, color_rgb=RGBColor(0, 0, 0))
    p_client.paragraph_format.space_after = Pt(12)

    add_cover_meta_field(doc, 'Project', header.get('project', 'Hitachi - Dynamics 365'))
    add_cover_meta_field(doc, 'Prepared by', header.get('prepared_by', ''))
    add_cover_meta_field(doc, 'Version', version, space_after=20)

    add_cover_table_heading(doc, 'Authored By:')
    add_cover_black_table(
        doc,
        ['Organization', 'Person'],
        [[authored_org, header.get('authored_person', '')]],
    )

    add_cover_table_heading(doc, 'Revision and signoff:')
    add_cover_black_table(
        doc,
        ['Date', 'Editor', 'Sections Revised', 'Description'],
        [[
            header.get('date', ''),
            header.get('revision_editor', ''),
            header.get('sections_revised', 'Initial'),
            header.get('revision_desc', 'Initiation'),
        ]],
        revision_data_row=True,
    )


def generate_tdd_docx(header, objects):
    doc = Document()

    # Set Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Hitachi Sans'
    font.size = Pt(10)

    # Page margins (extra top space for header logo + separator line)
    for section in doc.sections:
        section.top_margin = Inches(1.35)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_page_header(doc)
    add_cover_page(doc, header)

    # ── SECTION 1: Technical Design Planning ──
    add_section_heading(doc, '1.  Technical Design Planning', 1)
    add_section_heading(doc, 'Overview', 2)
    wi_num = header.get('work_item', '')
    p = doc.add_paragraph()
    run = p.add_run(f'This document is created to develop Enhancement {wi_num}: ')
    _style_run(run)
    run2 = p.add_run(header.get('enhancement_title', ''))
    _style_run(run2)

    # ── SECTION 2: Functional Requirement ──
    add_section_heading(doc, '2.  Functional Requirement', 1)
    p = doc.add_paragraph()
    run = p.add_run('Please refer VSTS for functional requirement')
    _style_run(run)

    # ── SECTION 3: Visual Studio Project ──
    add_section_heading(doc, '3.  Visual Studio Project', 1)
    add_bold_label(doc, 'Project Solutions:', header.get('project_solution', ''))

    # ── SECTION 4: User Interface ──
    forms = [o for o in objects if o['type'] == 'form']
    if forms:
        add_section_heading(doc, '4.  User Interface', 1)
        add_section_heading(doc, 'Forms', 2)

        for form in forms:
            add_hyperlink_heading(doc, form['name'])
            add_bold_label(doc, 'Description:', form.get('description', ''))

            props = form.get('properties', {})
            if any(props.values()):
                add_bold_label(doc, 'Form Details:')
                prop_table = doc.add_table(rows=1, cols=2)
                prop_table.style = 'Table Grid'
                prop_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                add_header_row(prop_table, ['Property', 'Value'])
                prop_map = [
                    ('Pattern', props.get('pattern', '')),
                    ('Style', props.get('style', '')),
                    ('Caption', props.get('caption', '')),
                    ('Data Source', props.get('data_source', '')),
                ]
                for label, val in prop_map:
                    if val:
                        add_data_row(prop_table, [label, val])
                doc.add_paragraph()

            controls_added = form.get('added_controls', [])
            if controls_added:
                add_bold_label(doc, 'Controls Added:')
                ctrl_table = doc.add_table(rows=1, cols=4)
                ctrl_table.style = 'Table Grid'
                add_header_row(ctrl_table, ['Control Name', 'Control Type', 'Data Source', 'Data Field'])
                for ctrl in controls_added:
                    add_data_row(ctrl_table, [
                        ctrl.get('name', ''),
                        ctrl.get('control_type', ''),
                        ctrl.get('data_source', ''),
                        ctrl.get('data_field', '')
                    ])
                doc.add_paragraph()

            controls_modified = form.get('modified_controls', [])
            if controls_modified:
                add_bold_label(doc, 'Modified Controls:')
                mod_table = doc.add_table(rows=1, cols=4)
                mod_table.style = 'Table Grid'
                add_header_row(mod_table, ['Control Name', 'Control Type', 'Data Source', 'Data Field'])
                for ctrl in controls_modified:
                    add_data_row(mod_table, [
                        ctrl.get('name', ''),
                        ctrl.get('control_type', ''),
                        ctrl.get('data_source', ''),
                        ctrl.get('data_field', '')
                    ])
                doc.add_paragraph()

            methods = form.get('methods', [])
            if methods:
                add_bold_label(doc, 'Methods:')
                meth_table = doc.add_table(rows=1, cols=2)
                meth_table.style = 'Table Grid'
                add_header_row(meth_table, ['Method Name', 'Description'])
                for m in methods:
                    add_data_row(meth_table, [m.get('name', ''), m.get('description', '')])
                doc.add_paragraph()

    # ── SECTION 5: Data Dictionary ──
    tables = [o for o in objects if o['type'] == 'table']
    views = [o for o in objects if o['type'] == 'view']
    edts = [o for o in objects if o['type'] == 'edt']
    enums = [o for o in objects if o['type'] == 'enum']
    data_entities = [o for o in objects if o['type'] == 'data_entity']
    
    if tables or views or edts or enums or data_entities:
        add_section_heading(doc, '5.  Data Dictionary', 1)

        if edts:
            add_section_heading(doc, 'EDT', 2)
            add_ax_subsection_table(
                doc,
                'EDT:',
                ['Name', 'Type', 'Extends'],
                [[e.get('name', ''), e.get('data_type', ''), e.get('extends', '')] for e in edts]
            )

        if enums:
            add_section_heading(doc, 'Base Enums', 2)
            add_ax_subsection_table(
                doc,
                'Enums:',
                ['Name', 'Description'],
                [[e.get('name', ''), e.get('description', '')] for e in enums]
            )

        if tables:
            add_section_heading(doc, 'Tables', 2)
            for tbl in tables:
                add_ax_table_section(doc, tbl)

        if views:
            add_section_heading(doc, 'Views', 2)
            for view in views:
                add_hyperlink_heading(doc, view['name'])
                add_bold_label(doc, 'Description:', view.get('description', ''))
                doc.add_paragraph()
        
        if data_entities:
            add_section_heading(doc, 'Data Entities', 2)
            for de in data_entities:
                add_hyperlink_heading(doc, de['name'])
                add_bold_label(doc, 'Description:', de.get('description', ''))
                doc.add_paragraph()

    # ── SECTION 6: Application Components ──
    classes = [o for o in objects if o['type'] == 'class']
    if classes:
        add_section_heading(doc, '6.  Application Components', 1)
        p = doc.add_paragraph()
        run = p.add_run('Following AX Classes/Menus/Action Menu Items created to handle the requirement.')
        _style_run(run)
        add_section_heading(doc, 'Classes', 2)

        for cls in classes:
            add_hyperlink_heading(doc, cls['name'])
            methods = cls.get('methods', [])
            if methods:
                m_table = doc.add_table(rows=1, cols=2)
                m_table.style = 'Table Grid'
                add_header_row(m_table, ['Method', 'Description'])
                for m in methods:
                    desc = m.get('description', '')
                    if m.get('is_new'):
                        desc = 'Added ' + desc if not desc.lower().startswith('added') else desc
                    add_data_row(m_table, [m.get('name', ''), desc])
                doc.add_paragraph()

    # ── SECTION 7: Services ──
    services_all = [o for o in objects if o['type'] == 'services']
    if services_all:
        add_section_heading(doc, '7.  Services', 1)
        
        service_groups = [s for s in services_all if s.get('subtype') == 'service_group']
        if service_groups:
            add_section_heading(doc, 'Service Groups', 2)
            sg_table = doc.add_table(rows=1, cols=2)
            sg_table.style = 'Table Grid'
            add_header_row(sg_table, ['Service Group Name', 'Service added under that service group'])
            for sg in service_groups:
                s_list = ', '.join([d.get('name', '') for d in sg.get('details', [])])
                add_data_row(sg_table, [sg.get('name', ''), s_list])
            doc.add_paragraph()

        services_only = [s for s in services_all if s.get('subtype') == 'service']
        if services_only:
            add_section_heading(doc, 'Services', 2)
            s_table = doc.add_table(rows=1, cols=2)
            s_table.style = 'Table Grid'
            add_header_row(s_table, ['Service Name', 'Methods'])
            for s in services_only:
                m_list = ', '.join([d.get('name', '') for d in s.get('details', [])])
                add_data_row(s_table, [s.get('name', ''), m_list])
            doc.add_paragraph()

    # ── SECTION 8: Security ──
    add_section_heading(doc, '8.  Security', 1)
    add_section_heading(doc, 'Security Privileges', 2)
    
    security_objects = [o for o in objects if o['type'] == 'security']
    if security_objects:
        for s in security_objects:
            add_hyperlink_heading(doc, f"{s.get('name', 'Security Object')} ({s.get('subtype', '')})")
            add_bold_label(doc, 'Description:', s.get('description', ''))
            
            if s.get('subtype') == 'privilege' and s.get('permissions'):
                p_table = doc.add_table(rows=1, cols=2)
                p_table.style = 'Table Grid'
                add_header_row(p_table, ['Object Name', 'Access Level'])
                for p in s.get('permissions', []):
                    add_data_row(p_table, [p.get('object_name', ''), p.get('access_level', '')])
                doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run('No security objects added yet')
        _style_run(run)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
